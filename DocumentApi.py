from flask import Flask, request, jsonify
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from werkzeug.utils import secure_filename
import os
import docx
from dotenv import load_dotenv
from docx import Document
from flask import Flask, request, send_file, jsonify
import gdown

load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'docx','doc'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


template = """Your task is to use the Main Input: {main_input}  and the four sample inputs provided which are
Sample Input 1: {sample_input_1}
Sample Input 2: {sample_input_2}
Sample Input 3: {sample_input_3}
Sample Input 4: {sample_input_4}

to analyze the information and propose the best possible input.
Be creative, thoughtful, and strategic in your approach. The quality of your proposed input 
will be based on how well you utilize the given data to craft an effective solution.

Best Input:"""

prompt = PromptTemplate(
    input_variables=["main_input", "sample_input_1", "sample_input_2", "sample_input_3", "sample_input_4"],
    template=template)

# Define the template for generating responses
template_topic = """
Legal Clause Generation

**Legal Topic:** {input_text}

As an experienced legal professional, your task is to provide legal clauses for the topic: "{input_text}." Your response should include typical legal clauses or content that can be incorporated into a legal document, such as a Non-Disclosure Agreement (NDA).

[Note: Each request for a specific legal topic will provide unique and relevant legal clauses.]
"""

# Create a PromptTemplate with only "input_text" as an input variable
prompt_topic = PromptTemplate(
    input_variables=["input_text"],  # Include only "input_text" here
    template=template_topic
)


def extract_paragraphs_from_docx(docx_path):
    paragraphs = []
    doc = docx.Document(docx_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        # Ignore empty paragraphs and paragraphs with only spaces
        if text:
            paragraphs.append(para)
    return paragraphs


def get_paragraph_style(para):
    para_style = para.style.name
    return para_style


def is_heading_style(para):
    heading_styles = ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']
    return any(style in get_paragraph_style(para) for style in heading_styles)


def get_key_from_paragraph(para):
    first_word = para.text.strip().split()[0]
    if is_heading_style(para):
        return para.text.strip().rstrip('.').rstrip(':')
    elif is_heading_style(para.runs[0]):
        return " ".join(para.text.strip().split()[:3]).rstrip('.').rstrip(':')
    return None


def create_topics_map(paragraphs):
    topics_map = {}
    current_topic = None
    related_paragraphs = []
    for para in paragraphs:
        key = get_key_from_paragraph(para)

        if key:
            if current_topic:
                current_topic = current_topic.strip().rstrip('.').rstrip(':')
                topics_map[current_topic] = related_paragraphs
                related_paragraphs = []

            current_topic = key

        if current_topic:
            related_paragraphs.append(para.text.strip())

    if current_topic is not None:
        current_topic = current_topic.strip().rstrip('.').rstrip(':')
        topics_map[current_topic] = related_paragraphs

    return topics_map


def download_file_from_google_drive(url, output_file):
    gdown.download(url, output_file, quiet=False)


llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_api_key, temperature=0)
finetune_chain = LLMChain(llm=llm, prompt=prompt)
finetune_chain_topic = LLMChain(llm=llm, prompt=prompt_topic)

@app.route('/generate_response', methods=['POST'])
def generate_response():
    data = request.get_json()
    input_text = data.get('input_text', '')

    # Generate a response using the model
    response = finetune_chain_topic.predict(input_text=input_text)

    return jsonify({'response': response})

@app.route('/find_paragraphs', methods=['POST'])
def search_topic_paragraphs():
    request_data = request.json
    topics = request_data.get('topics', [])
    file_path = request_data.get('docx_file_path', '')
    # Print the size of the topics list
    print(f"Size of topics: {len(topics)}")

    # Create a 'temp' directory in the current working directory
    temp_dir = os.path.join(os.getcwd(), 'temp')
    os.makedirs(temp_dir, exist_ok=True)
    filename = 'temp.docx'
    output_file = os.path.join(temp_dir, filename)
    print(output_file)
    download_file_from_google_drive(file_path, output_file)

    # Print the values in the topics list
    print("Values in topics:")
    for topic in topics:
        print(topic)
    if not output_file:
        return jsonify({"error": "Document file path not provided."}), 400

    # Load the document and create topics_map
    paragraphs = extract_paragraphs_from_docx(output_file)
    topics_map = create_topics_map(paragraphs)

    def spasifc(search_topic, doc_paragraphs):
        for i in range(len(doc_paragraphs)):
            if search_topic in doc_paragraphs[i].text:
                result_paragraphs = [doc_paragraphs[i].text]
                #for next_para in doc_paragraphs[i + 1:]:
                 #   if is_heading_style(next_para):
                  #      break
                   # result_paragraphs.append(next_para.text)
                return result_paragraphs
        return None

    result = {}
    for topic in topics:
        if topic in topics_map:
            # result[topic] = topics_map[topic]
            result[topic] = [value for value in topics_map[topic] if value != topic]
        else:
            print("Topic not found in the topics map. Searching in the document...")
            result_paragraphs = spasifc(topic, paragraphs)  # Note: here we use the original paragraphs variable

            if result_paragraphs:
                print(f"Found the text in the following paragraphs:")
                result[topic] = result_paragraphs
                for paragraph in result_paragraphs:
                    print(paragraph)
            else:
                print("Text not found in the document.")
                result[topic] = None

    return jsonify(result)


@app.route('/find/paragraphs', methods=['POST'])
def search_paragraphs():
    request_data = request.json
    topics = request_data.get('topics', [])
    file_path = request_data.get('docx_file_path', '')
    # Print the size of the topics list
    print(f"Size of topics: {len(topics)}")

    # Print the values in the topics list
    print("Values in topics:")
    for topic in topics:
        print(topic)
    if not file_path:
        return jsonify({"error": "Document file path not provided."}), 400

    # Load the document and create topics_map
    paragraphs = extract_paragraphs_from_docx(file_path)
    topics_map = create_topics_map(paragraphs)

    def spasifc(search_topic, doc_paragraphs):
        for i in range(len(doc_paragraphs)):
            if search_topic in doc_paragraphs[i].text:
                result_paragraphs = [doc_paragraphs[i].text]
                # for next_para in doc_paragraphs[i + 1:]:
                 #   if is_heading_style(next_para):
                  #      break
                   # result_paragraphs.append(next_para.text)
                return result_paragraphs
        return None

    result = {}
    for topic in topics:
        if topic in topics_map:
            # result[topic] = topics_map[topic]
            result[topic] = [value for value in topics_map[topic] if value != topic]
        else:
            print("Topic not found in the topics map. Searching in the document...")
            result_paragraphs = spasifc(topic, paragraphs)  # Note: here we use the original paragraphs variable

            if result_paragraphs:
                print(f"Found the text in the following paragraphs:")
                result[topic] = result_paragraphs
                for paragraph in result_paragraphs:
                    print(paragraph)
            else:
                print("Text not found in the document.")
                result[topic] = None

    return jsonify(result)


@app.route('/get_best_input', methods=['POST'])
def get_best_input():
    data = request.get_json()

    main_input = data.get('main_input', '')
    sample_input_1 = data.get('sample_input_1', '')
    sample_input_2 = data.get('sample_input_2', '')
    sample_input_3 = data.get('sample_input_3', '')
    sample_input_4 = data.get('sample_input_4', '')

    best_input = finetune_chain.predict(main_input=main_input, sample_input_1=sample_input_1,
                                        sample_input_2=sample_input_2, sample_input_3=sample_input_3,
                                        sample_input_4=sample_input_4)

    return jsonify({'best_input': best_input})


@app.route('/question-answering', methods=['POST'])
def question_answering():
    data = request.get_json()
    doc_path = data.get('doc_path')
    query = data.get('query')

    temp_dir = os.path.join(os.getcwd(), 'temp')
    os.makedirs(temp_dir, exist_ok=True)
    filename = 'temp.docx'
    output_file = os.path.join(temp_dir, filename)
    print(output_file)
    download_file_from_google_drive(doc_path, output_file)

    doc = Document(output_file)
    raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=800,
        chunk_overlap=10,
    )
    texts = text_splitter.split_text(raw_text)

    embeddings = OpenAIEmbeddings()
    docsearch = FAISS.from_texts(texts, embeddings)
    llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_api_key, temperature=0)
    chain = load_qa_chain(llm=llm, chain_type="stuff")

    docs = docsearch.similarity_search(query)
    answer = chain.run(input_documents=docs, question=query)

    return jsonify({'answer': answer})


@app.route('/question/answering', methods=['POST'])
def questionAnswering():
    data = request.get_json()
    doc_path = data.get('doc_path')
    query = data.get('query')


    doc = Document(doc_path)
    raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=800,
        chunk_overlap=10,
    )
    texts = text_splitter.split_text(raw_text)

    embeddings = OpenAIEmbeddings()
    docsearch = FAISS.from_texts(texts, embeddings)
    llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_api_key, temperature=0)
    chain = load_qa_chain(llm=llm, chain_type="stuff")

    docs = docsearch.similarity_search(query)
    answer = chain.run(input_documents=docs, question=query)

    return jsonify({'answer': answer})


@app.route('/upload', methods=['POST'])
def upload_file():
    # Check if the POST request has the file part
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    # If the user does not select a file, the browser submits an empty part without a filename
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        return jsonify({"file_path": file_path})
    return jsonify({"error": "File not allowed"}), 400

@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    file_path = os.path.join(filename)
    print("File Path:", file_path)  # Add this line for debugging
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"error": "File not found"}), 404

@app.route('/')
def hello_world():
    return 'welcome to flask app'


if __name__ == '__main__':
    app.run(debug=True, port=8080)

